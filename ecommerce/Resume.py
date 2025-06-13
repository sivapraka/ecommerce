from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a new Word document
doc = Document()

# Helper function to add a heading
def add_heading(text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# Helper function to add a normal paragraph with optional bold text
def add_paragraph(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.space_after = Pt(2)
    return p

# Add name and contact info center aligned
name = doc.add_paragraph()
name_run = name.add_run("Sivaprakash Rangasamy\n")
name_run.bold = True
name_run.font.size = Pt(16)
name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

contact_info = doc.add_paragraph("Senior Android Engineer | Jetpack Compose | Performance Optimization | Scalable Architecture\n"
                                 "Coimbatore, India • +91-8903964921 • sivaprakash.330570@gmail.com\n"
                                 "LinkedIn: linkedin.com/in/sivaprakash-rangasamy • LeetCode: leetcode.com/u/user0220y/")
contact_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
contact_info.space_after = Pt(12)

# Objective
add_heading("OBJECTIVE", level=2)
add_paragraph("Experienced Senior Android Engineer with expertise in Kotlin, Java, and modern Android frameworks. Proven track record in performance optimization and scalable architecture, complemented by backend skills in PHP (Laravel) and Python (Flask). Seeking to leverage a diverse technical background in a challenging, growth-oriented role.")

# Skills
add_heading("SKILLS", level=2)
skills_text = (
    "• Programming Languages: Kotlin, Java, PHP (Laravel), Python (Flask)\n"
    "• Android Frameworks: Jetpack Compose, Android Jetpack (Room, ViewModel, LiveData, Navigation), Retrofit, Hilt, WorkManager\n"
    "• Architectures: MVVM, Clean Architecture, Modularization\n"
    "• Performance Optimization: Memory Management, App Startup Time, Battery Efficiency, Profiling Tools\n"
    "• Databases: MySQL\n"
    "• Cloud Platforms: AWS\n"
    "• Testing & Tools: JUnit, Espresso, Mockito; Git, Firebase, CI/CD (GitHub Actions)"
)
add_paragraph(skills_text)

# Experience
add_heading("EXPERIENCE", level=2)

exp_list = [
    ("Senior Android Engineer – Ivy Mobility Solutions (Jan 2021 – Present)", [
        "Architected a modularized Android app, reducing build time by 35%.",
        "Migrated from XML layouts to Jetpack Compose, improving UI performance by 20%.",
        "Implemented an offline-first architecture using Room and WorkManager.",
        "Optimized app startup time by 40% via profiling tools and coroutine lazy loading.",
        "Led code reviews and mentorship programs, boosting team efficiency by 30%."
    ]),
    ("Senior Android Developer – Scoto Systec (Jan 2018 – Oct 2020)", [
        "Refactored a legacy codebase into an MVVM architecture, increasing maintainability by 50%.",
        "Optimized network calls with Retrofit and Coroutines, reducing API response time by 40%.",
        "Integrated Firebase Crashlytics to lower the crash rate by 30%."
    ]),
    ("Junior Android Developer – GNTS Technologies (Oct 2016 – Nov 2017)", [
        "Developed UI components using ConstraintLayout and ViewBinding.",
        "Implemented unit tests, increasing code coverage to 80%.",
        "Automated release pipelines with GitHub Actions, reducing deployment time by 60%."
    ])
]

for role, bullets in exp_list:
    add_paragraph(role, bold=True)
    for bullet in bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.space_after = Pt(0)

# Projects
add_heading("PROJECTS", level=2)

projects_list = [
    ("Bimbo – FMCG Delivery App", [
        "Developed a scalable Android application improving operational efficiency by 30%.",
        "Integrated Bluetooth thermal printing, reducing invoicing time by 25%.",
        "Implemented a multi-module architecture that accelerated build times.",
        "Secured app data using biometric authentication and encrypted shared preferences."
    ]),
    ("FMCG Solutions (B2B & B2C)", [
        "Created custom applications for the FMCG sector, increasing supply chain visibility and reducing order processing delays."
    ]),
    ("Healthcare Management System", [
        "Build doctor and patient apps for appointment scheduling.",
        "Designed DB schema, developed REST APIs in Laravel, and deployed backend on AWS."
    ]),
    ("IoT-Based Smart Systems", [
        "Car Parking Management: Integrated IoT sensors for real-time parking availability tracking.",
        "Student Management System: Developed a fingerprint-enabled attendance and academic tracking solution."
    ]),
    ("Textile Management", [
        "Reeling Management: Developed a system to manage silk reeling processes, ensuring operational accuracy.",
        "Attendance Tracking: Built an employee attendance module for shift-based workforce monitoring."
    ])
]

for project, bullets in projects_list:
    add_paragraph(project, bold=True)
    for bullet in bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.space_after = Pt(0)

# Education
add_heading("EDUCATION", level=2)
education_list = [
    ("Scaler Academy (Remote) – 2025 – Ongoing", "Specialized in Full Stack Development and Problem Solving."),
    ("Dr. NGP Institute of Technology – Coimbatore, India – 2016", "Bachelor of Engineering in Computer Science & Engineering.")
]

for edu, desc in education_list:
    add_paragraph(edu, bold=True)
    add_paragraph(desc)

if __name__ == '__main__':
    file_path = "D:/learning/Sivaprakash_Rangasamy_Resume_Formatted.docx"
    doc.save(file_path)
    print(doc.save(file_path))
